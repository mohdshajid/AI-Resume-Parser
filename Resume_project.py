import os
import json
from pathlib import Path

from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document


# ============================================================
# ENVIRONMENT
# ============================================================

load_dotenv()

api_key = os.getenv("GROQ_API_KEY")

if not api_key:
    raise ValueError("GROQ_API_KEY not found in .env")

client = Groq(api_key=api_key)

# You can change this in your .env if needed
MODEL = "openai/gpt-oss-120b"


# ============================================================
# JOB SCHEMA
# ============================================================

class JobD(BaseModel):
    role: str
    required_skills: list[str] = Field(default_factory=list)
    preferred_skills: list[str] = Field(default_factory=list)
    minimum_experience: float | None = None
    educational_requirements: list[str] = Field(default_factory=list)
    responsibilities: list[str] = Field(default_factory=list)


# ============================================================
# RESUME SCHEMA
# ============================================================

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = Field(default_factory=list)


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = Field(default_factory=list)
    experiences: list[Experience] = Field(default_factory=list)
    education: list[str] = Field(default_factory=list)
    projects: list[str] = Field(default_factory=list)
    certifications: list[str] = Field(default_factory=list)


# ============================================================
# MATCHING SCHEMA
# ============================================================

class MatchResult(BaseModel):
    score: float
    matching_skills: list[str] = Field(default_factory=list)
    missing_important_skills: list[str] = Field(default_factory=list)
    experience_requirement_met: bool
    final_verdict: str


# ============================================================
# PARSE JOB DESCRIPTION
# ============================================================

def parse_job(job_description: str) -> JobD:

    schema = JobD.model_json_schema()

    system_prompt = f"""
You are an expert HR hiring assistant.

Analyze the job description and extract structured information.

Return ONLY valid JSON matching this schema:

{schema}

Rules:
1. Do not return the schema itself.
2. Do not invent information.
3. If minimum experience is not mentioned, return null.
4. If a list has no information, return an empty list.
5. Extract the actual role, skills, education and responsibilities.
"""

    user_prompt = f"""
Analyze this job description:

{job_description}
"""

    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],
        response_format={"type": "json_object"}
    )

    data = json.loads(
        response.choices[0].message.content
    )

    return JobD(**data)


# ============================================================
# READ PDF
# ============================================================

def read_pdf(file_path: Path) -> str:

    reader = PdfReader(str(file_path))

    text = ""

    for page in reader.pages:

        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


# ============================================================
# READ DOCX
# ============================================================

def read_docx(file_path: Path) -> str:

    document = Document(str(file_path))

    text = ""

    for paragraph in document.paragraphs:

        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():
                    text += cell.text + "\n"

    return text


# ============================================================
# READ RESUME
# ============================================================

def read_resume(file_path: Path) -> str | None:

    extension = file_path.suffix.lower()

    if extension == ".pdf":

        return read_pdf(file_path)

    elif extension == ".docx":

        return read_docx(file_path)

    return None


# ============================================================
# PARSE RESUME
# ============================================================

def parse_resume(resume_text: str) -> Resume:

    schema = Resume.model_json_schema()

    system_prompt = f"""
You are an expert resume parser.

Extract information from the resume based on its meaning,
not only exact section headings.

For example:

Experience
Professional Experience
Work History
Employment
Internships

may all contain relevant experience.

Skills may appear in:
- Skills section
- Experience
- Internships
- Projects

Return ONLY valid JSON matching this schema:

{schema}

Rules:
1. Do not invent information.
2. If a value is unavailable, return null.
3. If a list has no information, return an empty list.
4. Include internships inside experiences.
5. Extract relevant skills from the entire resume.
6. Keep information concise and accurate.
"""

    user_prompt = f"""
Parse the following resume:

{resume_text}
"""

    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": user_prompt
            }
        ],
        response_format={"type": "json_object"}
    )

    data = json.loads(
        response.choices[0].message.content
    )

    return Resume(**data)


# ============================================================
# MATCH RESUME WITH JOB
# ============================================================

def match_candidate(
    job: JobD,
    resume: Resume
) -> MatchResult:

    schema = MatchResult.model_json_schema()

    prompt = f"""
You are an expert technical recruiter.

Compare the candidate resume against the job description.

JOB DESCRIPTION:

{job.model_dump_json(indent=2)}

CANDIDATE RESUME:

{resume.model_dump_json(indent=2)}

Return ONLY JSON matching this schema:

{schema}

Evaluate:

1. Matching skills
2. Missing important skills
3. Whether the experience requirement is met
4. Overall match score from 0 to 100
5. Final hiring verdict

Scoring guidance:

- 90-100 = Excellent match
- 75-89 = Strong match
- 60-74 = Moderate match
- 40-59 = Weak match
- 0-39 = Poor match

Do not invent candidate information.
Base the score on the actual resume and job description.
"""

    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        response_format={"type": "json_object"}
    )

    data = json.loads(
        response.choices[0].message.content
    )

    return MatchResult(**data)


# ============================================================
# RANK CANDIDATES
# ============================================================

def rank_candidates(candidates: list[dict]) -> list[dict]:

    return sorted(
        candidates,
        key=lambda candidate: candidate["match"].score,
        reverse=True
    )