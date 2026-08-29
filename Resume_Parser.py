import os
import time
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq 
from pydantic import BaseModel, Field

load_dotenv()
my_api_key=os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("API key not found")

client=Groq(api_key=my_api_key)

model="llama-3.3-70b-versatile"

job_description="""
Description
Do you want to solve real customer problems through innovative technology? Do you enjoy working on scalable services in a collaborative team environment? Do you want to see your code directly impact millions of customers worldwide?

At Amazon, we hire the best minds in technology to innovate and build on behalf of our customers. Customer obsession is part of our company DNA, which has made us one of the world's most beloved brands.

Our Software Development Engineers (SDEs) use modern technology to solve complex problems while seeing their work's impact first-hand. The challenges SDEs solve at Amazon are meaningful and influence millions of customers, sellers, and products globally. We seek individuals passionate about creating new products, features, and services while managing ambiguity in an environment where development cycles are measured in weeks, not years.

At Amazon, we believe in ownership at every level. As an SDE-I, you'll own the entire lifecycle of your code - from design through deployment and ongoing operations. This ownership mindset, combined with our commitment to operational excellence, ensures we deliver the highest quality solutions for our customers.

We're looking for curious minds who think big and want to define tomorrow's technology. At Amazon, you'll grow into the high-impact engineer you know you can be, supported by a culture of learning and mentorship. Every day brings exciting new challenges and opportunities for personal growth.

"""
class JobD(BaseModel):
    role:str
    required_skills:list[str]
    preferred_skills:list[str]
    minimum_experience:float|None
    educational_requirements:list[str]
    responsibilities:list[str]

jobd_schema=JobD.model_json_schema()

system_prompt=f"""
You are an expert hiring HR Assistant.
Your job is to analyze job description and extract structured information from them.
Return ONLY valid JSON matching this schema:
 
{jobd_schema}
IMPORTANT:
Do not return the schema itself.
Do not fields like "properties","title"or"type".
Fill the schema with actual information extracted from the job description.

If minimum experience is not mentioned, return null.
If information for a list is missing, return an empty list.
DO not INVENT information. 
 """

user_prompt=f"""
Analyze the following job descripiton:
{job_description}
"""
message_system={
    "role":"system",
    "content":system_prompt
}
message_user={
    "role":"user",
    "content":user_prompt
}
response_format={
    "type":"json_object"
}

messages=[message_system, message_user]

response=client.chat.completions.create(model=model, messages=messages, response_format=response_format)

answer=response.choices[0].message.content

raw_json=answer

#print (raw_json)

import json
job_data=json.loads(raw_json)
job=JobD(**job_data)
print(job.minimum_experience) 
print(job.educational_requirements)

#parse real
class MatchResult(BaseModel):
    score:float
    details:dict

class Experience(BaseModel):
        company:str|None=None
        role:str|None=None
        duration:str|None=None
        description:str|None=None
        skills_used:list[str]=[]

class Resume(BaseModel):
     name:str|None=None
     email:str|None=None
     phone:str|None=None
     total_experience_years:float|None=None
     skills:list[str]=[]
     experiences:list[Experience]=[]
     education:list[str]=[]
     projects:list[str]=[]
     certificatons:list[str]=[]

resume_schema=Resume.model_json_schema()
def final_score(job,resume):
    match_schema=MatchResult.model_json_schema()    
    prompt=f"""
    You are an HR recuiter.
    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}
    Give me:
    1.Candidate name
    2.Matching skills
    3.Missing important skills
    4.Whether experience requirement is met
    5.Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.    
    """  
    message={
     "role":"user",
     "content":prompt
    }
    messages=[message]
    response_format={
     "type":"json_object"
}
    response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)
    data=json.loads(response.choices[0].message.content)
    return MatchResult(**data)

def parse_resume(resume_text):
     system_prompt=f"""
    You are an expert resume parser.
    Extract information from the resume based on its meaning,
    not only based on exact section haeding.

    Different resumes may use different heaings.
    For example:
    -Experience
    -Professional Experience
    -Work History
    -Employement
    -Internships

    These may all contain relevant experience.
    Skills may also appear in the skills section,work experience,internship or projects.
    return ONLY valid JSON matching this schema:

    {resume_schema}
    Important rules:
    1.Do not invent information.
    2.If a value is not available,return null.
    3.If a list has no information,return an empty list
    4.Include internships inside experiences.
    5.Extract skills mentioned across the entire resume.
    """
     user_prompt=f"""
    Parse the following resume:
    {resume_text}
    """
     message_system={
          "role":"system",
          "content":system_prompt
     }
     message_user={
          "role":"user",
          "content":user_prompt
     }
     messages=[message_system,message_user]
     response_format={
          "type":"json_object"
     }
     response=client.chat.completions.create(model=model,messages=messages,response_format=response_format)
     raw_output=response.choices[0].message.content
     data=json.loads(raw_output)
     resume=Resume(**data)
     return resume

from pypdf import PdfReader
from docx import Document
def read_pdf(file_path):
     reader=PdfReader(file_path)
     text=""
     for page in reader.pages:
          page_text=page.extract_text()
          if page_text:
               text += page_text + "\n"
     return text  

def read_docx(file_path):
     document=Document(file_path)
     text=""
     for paragraph in document.paragraphs:
          if paragraph.text.strip():
               text += paragraph.text + "\n"

     for table in document.tables:
          for row in table.rows:
               for cell in row.cells:
                    if cell.text.strip():
                         text += cell.text + "\n"
     return text 

def read_resume(file_path):
     if file_path.suffix.lower()==".pdf":
          return read_pdf(file_path)
     elif file_path.suffix.lower()==".docx":
          return read_docx(file_path)
     else:
          return None


resume_folder=Path("resumes")
all_results=[]
for file_path in resume_folder.iterdir():
     if file_path.suffix.lower()not in [".pdf",".docx"]:
          continue
     print ("/nProcessing:",file_path.name)
     resume_text=read_resume(file_path)
     parsed_resume=parse_resume(resume_text)
     time.sleep(5)
     result=final_score(job,parsed_resume)
     time.sleep(5)
     print("Score:",result.score)
     all_results.append({
          "name":parsed_resume.name,
          "score":result.score,
          "details":result.details
     })
     all_results.sort(
          key=lambda candidate:candidate["score"],
          reverse=True
     )
     top_2=all_results[:2]

     print("TOP 2 CANDIDATES")
     for candidate in top_2:

          print(
               candidate["name"],
               "_",
               candidate["score"],
               "_",
          )
          print (candidate["details"])
       

     
                                            
             



